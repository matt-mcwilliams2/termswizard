import re
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(agreement_text: str) -> bytes:
    doc = Document()

    # Set normal margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = agreement_text.strip().split("\n")
    title = lines[0].strip() if lines else "Affiliate Terms & Conditions"

    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(24)

    # Process remaining lines
    in_numbered_list = False
    for line in lines[1:]:
        stripped = line.strip()

        # Empty line
        if not stripped:
            in_numbered_list = False
            continue

        # Section headers (ALL CAPS, not numbered, not a bullet)
        if (
            stripped == stripped.upper()
            and len(stripped) > 3
            and not stripped[0].isdigit()
            and not stripped.startswith("-")
        ):
            in_numbered_list = False
            para = doc.add_paragraph()
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(16)
            continue

        # Numbered items
        numbered_match = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if numbered_match:
            in_numbered_list = True
            para = doc.add_paragraph(stripped, style="List Number")
            for run in para.runs:
                run.font.size = Pt(12)
            continue

        # Bullet items
        if stripped.startswith("- "):
            in_numbered_list = False
            para = doc.add_paragraph(stripped[2:], style="List Bullet")
            for run in para.runs:
                run.font.size = Pt(12)
            continue

        # Regular paragraph
        in_numbered_list = False
        para = doc.add_paragraph()
        run = para.add_run(stripped)
        run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
